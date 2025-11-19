# src/resume_processor.py
import os
import PyPDF2
import docx
from src.nlp_utils import tokenize, calculate_similarity
import re
import streamlit as st

import pandas as pd

# import pandas as pd
# import os
# import PyPDF2
# import docx
# import re


# def load_candidate_data():
#     # Load candidate data from a CSV file or any other source
#     # For demonstration, we will create a sample DataFrame
#     data = {
#         'Name': ['Alice', 'Bob', 'Charlie', 'David', 'Eve'],
#         'Skills': [
#             'Python, Machine Learning, Data Analysis',
#             'Java, Spring, Hibernate',
#             'Python, Data Science, Machine Learning',
#             'JavaScript, React, Node.js',
#             'Python, Flask, Data Visualization'
#         ],
#         'Experience': [
#             '3 years in software development',
#             '5 years in backend development',
#             '4 years in data science',
#             '2 years in frontend development',
#             '3 years in web development'
#         ]
#     }
#     return pd.DataFrame(data)

# ++++++++++++++++++++++++++++++++++++++++++++++++++


def extract_text_from_pdf(file_path):
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

import re

# src/data_preparation.py
import re

# src/data_preparation.py
import re

def parse_resume_text(resume_text, filename):
    # Simple parsing logic to extract name, skills, and qualifications
    name_pattern = r'(?i)Name:\s*(.*?)(?=\n|$)'  # Case insensitive match for "Name:"
    skills_pattern = r'(?i)Skills:\s*(.*?)(?=\n|$)'  # Case insensitive match for "Skills:"
    qualifications_pattern = r'(?i)Qualifications:\s*(.*?)(?=\n|$)'  # New section for qualifications

    name = re.search(name_pattern, resume_text)
    skills = re.search(skills_pattern, resume_text)
    qualifications = re.search(qualifications_pattern, resume_text)

    return {
        'Name': name.group(1).strip() if name else filename,  # Use filename if name is not found
        'Skills': skills.group(1).strip() if skills else "N/A",
        'Qualifications': qualifications.group(1).strip() if qualifications else "N/A"  # New section
    }

def load_candidate_data_from_resumes(uploaded_files):
    candidates = []
    
    for uploaded_file in uploaded_files:
        file_path = os.path.join("data", uploaded_file.name)
        
        # Save the uploaded file temporarily
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        # Extract text based on file type
        if uploaded_file.type == "application/pdf":
            resume_text = extract_text_from_pdf(file_path)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            resume_text = extract_text_from_docx(file_path)
        else:
            continue  # Skip unsupported file types
        
        # Parse the resume text to extract candidate information
        candidate_info = parse_resume_text(resume_text, uploaded_file.name)  # Pass the filename
        candidates.append(candidate_info)
    
    # Create a DataFrame from the list of candidates
    return pd.DataFrame(candidates)



#+++++++++++++++++++++++++++++++++++++++++++++++++++


def extract_text_from_pdf(file_path):
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def extract_contact_info(resume_text):
    # Regular expressions for extracting email and phone numbers
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    phone_pattern = r'\+?\d[\d -]{8,12}\d'

    emails = re.findall(email_pattern, resume_text)
    phones = re.findall(phone_pattern, resume_text)

    return emails, phones

def process_resumes(uploaded_files, job_description, skills, experience):
    if not uploaded_files:
        st.error("Please upload at least one resume.")
    if not job_description:
        st.error("Please provide a job description.")
    
    resume_data = {}
    
    # Extract text from each uploaded resume
    for uploaded_file in uploaded_files:
        file_path = os.path.join("data", uploaded_file.name)
        if uploaded_file.type == "application/pdf":
            resume_text = extract_text_from_pdf(file_path)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            resume_text = extract_text_from_docx(file_path)
        else:
            continue  # Skip unsupported file types

        # Extract contact information
        emails, phones = extract_contact_info(resume_text)

        # Tokenize the resume text
        tokens = tokenize(resume_text)
        resume_data[uploaded_file.name] = {
            "text": resume_text,
            "tokens": tokens,
            "emails": emails,
            "phones": phones
        }

    # Rank resumes based on job description and skills
    rankings = rank_resumes(resume_data, job_description, skills, experience)
    return rankings, resume_data

def rank_resumes(resume_data, job_description, skills, experience):
    rankings = {}
    
    # Tokenize job description and skills
    job_tokens = tokenize(job_description)
    skill_tokens = tokenize(skills)

    for resume_name, data in resume_data.items():
        # Calculate similarity score
        similarity_score = calculate_similarity(data["text"], job_description)
        rankings[resume_name] = similarity_score

    # Sort rankings based on similarity score
    sorted_rankings = dict(sorted(rankings.items(), key=lambda item: item[1], reverse=True))
    return sorted_rankings